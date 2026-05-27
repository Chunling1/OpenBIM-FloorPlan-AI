"""
生成 M2 DA 训练报告 Word 文档（含所有图片）
"""
import os
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).parent
REPORT_DIR = BASE_DIR / "output_paper" / "report"
OUTPUT_DOCX = BASE_DIR / "output_paper" / "M2_DA_训练报告.docx"


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None, header_color='2C3E50'):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        set_cell_shading(cell, header_color)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'F8F9FA')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def main():
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 标题页 =====
    doc.add_paragraph()  # 空行
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('M2 UNet-ResNet34\nDomain Adaptation 训练报告')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('户型图语义分割 · 域适应增强')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(127, 140, 141)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        '训练完成时间: 2026-04-15 05:19\n'
        '总训练时长: 926.4 分钟 (~15.4 小时)\n'
        'GPU: NVIDIA GeForce RTX 4050 Laptop (6GB VRAM)\n'
        '模型参数: 24,436,804 (24.4M)'
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ===== 目录 =====
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 训练概述',
        '2. 训练曲线与指标分析',
        '3. 验证集效果对比 (Ground Truth vs Prediction)',
        '4. 真实CAD图推理效果对比',
        '5. 所有模型横向对比',
        '6. 结论与下一步',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_page_break()

    # ===== 1. 训练概述 =====
    doc.add_heading('1. 训练概述', level=1)

    doc.add_heading('1.1 模型架构', level=2)
    arch_items = [
        ('模型', 'U-Net + ResNet34 encoder (ImageNet pretrained)'),
        ('参数量', '24,436,804 (24.4M)'),
        ('输入尺寸', '512 × 512'),
        ('输出类别', '4类: Background, Wall, Window, Door'),
    ]
    for k, v in arch_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{k}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(v)
        run.font.size = Pt(10)

    doc.add_heading('1.2 训练配置', level=2)
    add_styled_table(doc,
        ['配置项', '值'],
        [
            ['Batch Size', '2 (× 4 grad accum = effective 8)'],
            ['Optimizer', 'AdamW'],
            ['Encoder LR', '3e-5 (差分学习率)'],
            ['Decoder LR', '3e-4'],
            ['Weight Decay', '1e-4'],
            ['Scheduler', 'LinearWarmup(5ep) → CosineAnnealing'],
            ['Loss', 'Weighted CrossEntropy + Dice Loss'],
            ['Class Weights', '[0.5, 2.0, 3.0, 3.0]'],
            ['Gradient Clipping', '0.5'],
            ['Precision', 'FP32 (AMP disabled for stability)'],
            ['Dataset', 'CubiCasa5K (4200 train / 400 val)'],
            ['Epochs', '100'],
        ],
        col_widths=[5, 10]
    )

    doc.add_paragraph()
    doc.add_heading('1.3 Domain Adaptation 增强策略', level=2)
    p = doc.add_paragraph(
        '为使模型泛化到黑底中国CAD图，在标准数据增强的基础上，添加了以下域适应增强：'
    )

    add_styled_table(doc,
        ['增强方式', '概率', '目的'],
        [
            ['InvertImg', 'p=0.3', '随机反色，模拟黑底CAD图'],
            ['ToGray', 'p=0.2', '转灰度，模拟CAD线稿'],
            ['ColorJitter', 'p=0.3', '大幅色彩抖动(B=0.4,C=0.4,S=0.3,H=0.1)'],
        ],
        col_widths=[4, 3, 8]
    )

    p = doc.add_paragraph()
    run = p.add_run('标准增强还包括: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(
        'RandomResizedCrop, HorizontalFlip(0.5), VerticalFlip(0.3), '
        'RandomRotate90(0.3), ElasticTransform/GridDistortion(0.15), '
        'GaussNoise/GaussianBlur(0.15), RandomBrightnessContrast(0.3), '
        'CoarseDropout(0.15)'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ===== 2. 训练曲线 =====
    doc.add_heading('2. 训练曲线与指标分析', level=1)

    img_path = REPORT_DIR / "01_training_curves.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(6.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading('2.1 关键观察', level=2)

    observations = [
        ('Loss 曲线', 'Train/Val loss 平稳收敛，无震荡或崩溃，训练非常稳定。Train loss 从 1.55 降到 0.28，Val loss 从 1.12 降到 0.30。'),
        ('mIoU 曲线', '前20 epoch快速上升（0.33→0.73），20-60 epoch稳步提升，60+ epoch进入平台期。最终 Train mIoU 0.790, Val mIoU 0.778。'),
        ('各类别 IoU', 'Background 始终最高(~0.97)，Wall 和 Window 稳步提升至 0.73/0.77，Door 是最大短板(0.66)。'),
        ('无灾难性遗忘', '相比之前的微调尝试，DA augmentation 策略成功避免了性能崩溃。'),
    ]
    for title, detail in observations:
        p = doc.add_paragraph()
        run = p.add_run(f'• {title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(detail)
        run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_heading('2.2 最终训练指标', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Best Model @ Epoch 78')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(39, 174, 96)

    add_styled_table(doc,
        ['指标', '值'],
        [
            ['Best Val mIoU', '0.7821'],
            ['Background IoU', '0.9693'],
            ['Wall IoU', '0.7358'],
            ['Window IoU', '0.7656'],
            ['Door IoU', '0.6577'],
        ],
        col_widths=[5, 5]
    )

    doc.add_paragraph()

    add_styled_table(doc,
        ['指标', 'Train (E100)', 'Val (E100)'],
        [
            ['Loss', '0.2838', '0.3029'],
            ['mIoU', '0.790', '0.778'],
        ],
        col_widths=[4, 4, 4]
    )

    doc.add_page_break()

    # ===== 3. 验证集 GT vs Prediction =====
    doc.add_heading('3. 验证集效果对比', level=1)

    p = doc.add_paragraph(
        '从验证集(400样本)中均匀选取8个样本，展示 原图 → Ground Truth Mask → M2-DA 模型预测 三列对比。'
        '颜色编码: 红色=Wall, 蓝色=Window, 绿色=Door, 深灰=Background。'
    )
    for run in p.runs:
        run.font.size = Pt(10)

    img_path = REPORT_DIR / "02_val_gt_vs_pred.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(6.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading('3.1 分析', level=2)

    analysis = [
        ('墙体 (Wall)', '预测质量最好，主要轮廓基本准确，与GT高度吻合。偶有过分割（将走廊/过道误判为墙体）。'),
        ('窗户 (Window)', '在多数样本中定位准确，但偶有漏检（尤其是小窗），部分窗被误分为墙。'),
        ('门 (Door)', '最具挑战性：小门容易被忽略，与墙体交界处边界模糊，部分被错误预测为窗或墙。'),
    ]
    for title, detail in analysis:
        p = doc.add_paragraph()
        run = p.add_run(f'• {title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(detail)
        run.font.size = Pt(10)

    doc.add_page_break()

    # ===== 4. CAD 图对比 =====
    doc.add_heading('4. 真实CAD图推理效果对比', level=1)

    p = doc.add_paragraph(
        '使用4张中国黑底CAD图进行推理，对比3种推理策略: '
        '① DA模型直接推理(黑底原图)  ② 原版模型+预处理(反色→白底)  ③ DA模型+预处理。'
    )
    for run in p.runs:
        run.font.size = Pt(10)

    img_path = REPORT_DIR / "03_cad_da_vs_orig.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(6.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading('4.1 策略对比', level=2)

    add_styled_table(doc,
        ['策略', 'Wall', 'Window', 'Door', '评价'],
        [
            ['DA + Direct (黑底原图)', '★★★', '★★', '★', '能直接识别黑底图，Wall检出好'],
            ['Orig + Preprocess (白底转换)', '★★★★', '★★', '★', '预处理转白底后效果提升明显'],
            ['DA + Preprocess (最佳组合)', '★★★★', '★★★', '★★', 'DA模型+预处理叠加效果最好'],
        ],
        col_widths=[5, 2, 2, 2, 5]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('推荐生产策略: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(39, 174, 96)
    run = p.add_run('DA模型 + 黑底预处理（反色+CLAHE增强），可获得最佳泛化效果。')
    run.font.size = Pt(11)

    doc.add_page_break()

    # ===== 5. 模型对比 =====
    doc.add_heading('5. 所有模型横向对比', level=1)

    img_path = REPORT_DIR / "04_model_comparison.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(6.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading('5.1 性能排名', level=2)

    add_styled_table(doc,
        ['排名', '模型', 'Best mIoU', '参数量'],
        [
            ['1', 'M2: UNet+ResNet34 (原版)', '0.7861', '24.4M'],
            ['2', 'M2-DA: UNet+ResNet34+DA', '0.7821', '24.4M'],
            ['3', 'M3: DeepLabV3++EfficientNet-B4', '0.7615', '18.5M'],
            ['4', 'M1: LightUNet (baseline)', '0.6847', '7.8M'],
        ],
        col_widths=[2, 7, 3, 3]
    )

    doc.add_paragraph()
    doc.add_heading('5.2 分析', level=2)

    comparisons = [
        'M2-DA 在验证集 mIoU (0.7821) 仅比无DA的M2 (0.7861) 低 0.004，几乎无损。',
        'M2-DA 额外获得了对黑底CAD图的泛化能力，这是无DA版本完全不具备的。',
        'M1 (LightUNet) 训练过程中出现严重不稳定（训练崩溃+恢复），最终mIoU只有0.685。',
        'M3 表现不如M2，可能与EfficientNet在这个规模数据上的适配性有关。',
    ]
    for c in comparisons:
        p = doc.add_paragraph(c, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ===== 6. 结论 =====
    doc.add_heading('6. 结论与下一步', level=1)

    doc.add_heading('6.1 成功点', level=2)
    successes = [
        ('DA策略有效', '仅靠数据增强（InvertImg + ToGray + ColorJitter）就让模型获得了域泛化能力，无需额外标注数据。'),
        ('训练稳定', '100 epoch 全程无崩溃、无灾难性遗忘，Cosine Annealing 调度工作良好。'),
        ('性能损失极小', '验证集 mIoU 仅下降 0.4%（0.7861 → 0.7821），代价完全可接受。'),
    ]
    for title, detail in successes:
        p = doc.add_paragraph()
        run = p.add_run(f'✅ {title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(39, 174, 96)
        run = p.add_run(detail)
        run.font.size = Pt(10)

    doc.add_heading('6.2 待改进', level=2)
    improvements = [
        ('Door IoU 偏低 (0.658)', '可考虑增大Door的类别权重（当前3.0→4.0），或加入更多门的正样本。'),
        ('CAD图窗户检出率不足', '中国CAD图中窗户的表达方式（多线段/填充图案）与CubiCasa数据集差异大，需要更针对性的增强或少量标注数据。'),
        ('推理管线集成', '应将预处理步骤（反色+CLAHE）集成到生产推理管线中，确保最佳效果。'),
    ]
    for title, detail in improvements:
        p = doc.add_paragraph()
        run = p.add_run(f'⚠️ {title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(243, 156, 18)
        run = p.add_run(detail)
        run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_heading('6.3 文件位置', level=2)
    files = [
        ('最佳模型', 'models/M2_UNet_ResNet34_DA_best.pt'),
        ('训练日志', 'm2_da_train.log'),
        ('训练历史', 'output_paper/M2_UNet_ResNet34_DA_history.json'),
        ('报告图片', 'output_paper/report/'),
    ]
    for k, v in files:
        p = doc.add_paragraph()
        run = p.add_run(f'{k}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(v)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(41, 128, 185)

    # 保存
    doc.save(str(OUTPUT_DOCX))
    print(f"Word report saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
